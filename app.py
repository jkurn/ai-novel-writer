import streamlit as st
import openai
from docx import Document
import os
import json
from dotenv import load_dotenv
from io import BytesIO

# Load environment variables from .env file
load_dotenv()

# Initialize OpenAI API (you'll need to set up your API key)
api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    raise ValueError("The OPENAI_API_KEY environment variable is not set.")
client = openai.OpenAI(api_key=api_key)

NOVEL_DIR = "novels"

def generate_chapter(market_understanding, style_guide, story_outlook, min_words):
    messages = [
        {"role": "system", "content": "You are an author for lightnovel content"},
        {"role": "user", "content": f"""
        You are tasked with developing ideas for a 100+ chapter webnovel to be published on Fizzo. To accomplish this, you will be provided with market understanding, style guide requirements, and a story outlook. Your goal is to analyze this information and generate compelling ideas for a long-form webnovel.

        First, review the market understanding:
        <market_understanding>
        {market_understanding}
        </market_understanding>

        Next, familiarize yourself with the style guide requirements:
        <style_guide>
        {style_guide}
        </style_guide>

        Now, consider the story outlook provided:
        <story_outlook>
        {story_outlook}
        </story_outlook>

        Analyze the information provided above, paying close attention to:
        1. Target audience preferences and trends
        2. Popular themes and genres in the market
        3. Specific style and formatting requirements
        4. The overall direction and tone suggested by the story outlook

        Using this analysis, brainstorm ideas for a 100+ chapter webnovel that aligns with the market understanding, adheres to the style guide, and builds upon the story outlook. Consider the following elements:

        1. Main plot arc
        2. Subplots that can span multiple chapters
        3. Character development opportunities
        4. World-building elements
        5. Potential plot twists and cliffhangers
        6. Themes and motifs that can be explored throughout the novel

        Organize your thoughts in a structured format, using <brainstorming> tags. Within these tags, use subheadings for different aspects of your webnovel concept.

        Once you have completed your brainstorming, provide a final output in the following format:

        <webnovel_concept>
        <title>Proposed title for the webnovel</title>
        <genre>Primary and secondary genres</genre>
        <summary>A brief 2-3 sentence summary of the overall story</summary>
        <main_characters>List 3-5 main characters with brief descriptions</main_characters>
        <plot_outline>
        Provide a high-level outline of the main plot, divided into 3-5 major arcs. Each arc should have potential for 20-30 chapters.
        </plot_outline>
        <key_themes>List 3-5 key themes that will be explored</key_themes>
        </webnovel_concept>
        """}
    ]
    
    # Create an empty container
    container = st.empty()
    
    try:
        chapter_content = ""
        for response in client.chat.completions.create(
            model="gpt-4o-mini",  # Make sure this is the correct model name
            messages=messages,
            temperature=1,
            max_tokens=min_words * 5,  # Adjust as needed
            stream=True  # Enable streaming
        ):
            if hasattr(response.choices[0].delta, 'content'):
                chunk_message = response.choices[0].delta.content
                if chunk_message is not None:
                    chapter_content += chunk_message
                    # Update the container with the accumulated content
                    container.markdown(chapter_content)
            
        if not chapter_content:
            st.error("No content was generated. Please check your API key and model settings.")
            return None

        return chapter_content.strip()
    except Exception as e:
        st.error(f"An error occurred while generating the chapter: {str(e)}")
        return None

def get_docx(novel_title, chapter_number, chapter_title, content):
    doc = Document()
    doc.add_heading(f"{novel_title}", 0)
    doc.add_heading(f"Chapter {chapter_number}: {chapter_title}", 1)
    doc.add_paragraph(content)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def save_novel(novel, filename):
    with open(filename, 'w') as f:
        json.dump(novel, f, indent=2)

def load_novel(filename):
    with open(filename, 'r') as f:
        return json.load(f)

def main():
    st.title("AI Book Writer Continuation")

    # Initialize session state for novel
    if 'novel' not in st.session_state:
        st.session_state.novel = None

    # Novel management
    st.sidebar.header("Novel Management")
    novel_title = st.sidebar.text_input("Novel Title")
    novel_filename = os.path.join(NOVEL_DIR, f"{novel_title}.json")

    if st.sidebar.button("Create New Novel"):
        if novel_title:
            st.session_state.novel = {"title": novel_title, "chapters": []}
            save_novel(st.session_state.novel, novel_filename)
            st.sidebar.success(f"New novel '{novel_title}' created")
        else:
            st.sidebar.error("Please enter a novel title")

    if st.sidebar.button("Load Existing Novel"):
        if os.path.exists(novel_filename):
            st.session_state.novel = load_novel(novel_filename)
            st.sidebar.success(f"Novel '{novel_title}' loaded")
        else:
            st.sidebar.error(f"Novel '{novel_title}' does not exist")

    if st.session_state.novel:
        st.sidebar.write(f"Current Novel: {st.session_state.novel['title']}")
        if st.sidebar.button("Save Current Novel"):
            save_novel(st.session_state.novel, novel_filename)
            st.sidebar.success(f"Novel '{st.session_state.novel['title']}' saved")

        # Add chapter number input
        chapter_number = st.number_input("Chapter Number", min_value=1, value=len(st.session_state.novel['chapters']) + 1)

        # User inputs for chapter generation
        market_understanding = st.text_area("Market Understanding", value=open("market.txt").read())
        style_guide = st.text_area("Style Guide", value=open("style.txt").read())
        story_outlook = st.text_area("Story Outlook")
        min_words = st.number_input("Minimum Word Count", min_value=100, value=500)

        if st.button("Generate Chapter"):
            if market_understanding and style_guide and story_outlook:
                with st.spinner("Generating chapter..."):
                    chapter_content = generate_chapter(market_understanding, style_guide, story_outlook, min_words)
                if chapter_content:
                    chapter_title = chapter_content.split('\n')[0]  # Assume the first line is the title
                    st.write(f"Generated Chapter {chapter_number}: {chapter_title}")
                    st.write(chapter_content)

                    # Add chapter to novel at the specified index
                    new_chapter = {
                        "number": chapter_number,
                        "title": chapter_title,
                        "content": chapter_content
                    }
                    if chapter_number > len(st.session_state.novel['chapters']):
                        st.session_state.novel['chapters'].append(new_chapter)
                    else:
                        st.session_state.novel['chapters'].insert(chapter_number - 1, new_chapter)
                    save_novel(st.session_state.novel, novel_filename)

                    # Store the generated chapter in session state
                    st.session_state.last_generated_chapter = new_chapter
                else:
                    st.error("Failed to generate chapter. Please check the console for more information.")
            else:
                st.error("Please fill in all fields")

        if 'last_generated_chapter' in st.session_state:
            chapter = st.session_state.last_generated_chapter
            if isinstance(chapter, str):
                # Old format: chapter is just a string of content
                chapter_number = len(st.session_state.novel['chapters'])
                chapter_title = chapter.split('\n')[0]  # Assume first line is title
                chapter_content = chapter
            else:
                # New format: chapter is a dictionary
                chapter_number = chapter['number']
                chapter_title = chapter['title']
                chapter_content = chapter['content']

            st.download_button(
                "Download Chapter as Word Document",
                data=get_docx(st.session_state.novel['title'], chapter_number, chapter_title, chapter_content),
                file_name=f"chapter_{chapter_number}_{chapter_title}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    if not os.path.exists(NOVEL_DIR):
        os.makedirs(NOVEL_DIR)
    main()